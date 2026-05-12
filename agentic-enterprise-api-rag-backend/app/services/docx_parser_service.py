from __future__ import annotations

import re
from collections.abc import Iterator
from io import BytesIO
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

# Legacy templates used DSE|NOK|SAC; QR Util and others use API-REST-<LETTERS>-<digits> (e.g. API-REST-QRU-01).
API_REF_PATTERN = re.compile(r"API-REST-[A-Z]+-\d+", re.IGNORECASE)
_GET_SSO_TOKEN_HEADING = re.compile(r"^getSSOToken\b", re.IGNORECASE)
_QR_CHECKOUT_HEADING = re.compile(
    r"QR\s*Checkout|generate-qr|/v1/provisioning|provisioning/generate",
    re.IGNORECASE,
)
# Caption lines immediately before a parameter table (QR Util and similar templates).
_RESPONSE_TABLE_CAPTION_RE = re.compile(
    r"^\s*(response\s*parameters|output\s*response(?:\s*success)?|success\s*response|output\s*response\s*success)\s*$",
    re.IGNORECASE,
)
_REQUEST_TABLE_CAPTION_RE = re.compile(
    r"^\s*(request\s*parameters|header\s*parameters|query\s*parameters|body\s*parameters|input\s*parameter(?:s)?)\s*$",
    re.IGNORECASE,
)
_ERROR_TABLE_CAPTION_RE = re.compile(
    r"^\s*(expected\s*error\s*codes?|error\s*codes?)\s*$",
    re.IGNORECASE,
)
_JWT_TABLE_CAPTION_RE = re.compile(
    r"^\s*(jwt\s*payload|payload\s*structure|token\s*payload)\s*$",
    re.IGNORECASE,
)


class DocxParserService:
    def parse_preview(self, file_name: str, file_bytes: bytes) -> dict[str, Any]:
        document = Document(BytesIO(file_bytes))
        logical_sections = self._extract_logical_sections(document)
        apis = self._parse_document(document, logical_sections)
        return {
            "file_name": file_name,
            "document_title": self._extract_document_title(document),
            "purpose_scope": self._extract_purpose_scope(document),
            "api_count": len(apis),
            "logical_sections": logical_sections,
            "authentication_preamble": self._extract_authentication_preamble(logical_sections),
            "apis": [
                {
                    "api_reference_id": api["api_reference_id"],
                    "service_name": api.get("service_name"),
                    "service_group": api.get("service_group"),
                    "service_method": api.get("service_method"),
                    "service_pattern": api.get("service_pattern"),
                }
                for api in apis
            ],
            "apis_full": apis,
        }

    def _extract_authentication_preamble(self, logical_sections: list[dict[str, str]]) -> str:
        """General OAuth / authentication narrative (not tied to a single REST id)."""
        parts: list[str] = []
        for section in logical_sections:
            title = (section.get("section_title") or "").strip()
            tl = title.lower()
            content = self._normalize_text(section.get("content") or "")
            if not content:
                continue
            if (
                "general authentication" in tl
                or "oauth2 client credentials" in tl
                or tl == "oauth2"
                or ("oauth" in tl and "credential" in tl)
                or ("client credentials" in tl and "grant" in content.lower())
                or ("access token" in tl and "expir" in content.lower())
                or "refresh token" in tl
                or "getssotoken" in content.lower()
            ):
                parts.append(f"{title}\n{content}")
        return "\n\n".join(parts).strip()

    def _extract_logical_sections(self, document: DocumentType) -> list[dict[str, str]]:
        collected: list[dict[str, str]] = []
        current_title = "General"
        current_lines: list[str] = []
        for block in self._iter_block_items(document):
            if isinstance(block, Paragraph):
                text = self._normalize_text(block.text)
                if not text:
                    continue
                if self._is_heading_like(block, text):
                    if current_lines:
                        collected.append(
                            {
                                "section_title": current_title,
                                "content": "\n".join(current_lines).strip(),
                            }
                        )
                        current_lines = []
                    current_title = text
                    continue
                current_lines.append(text)
            elif isinstance(block, Table):
                table_rows = self._extract_table_rows(block)
                if not table_rows:
                    continue
                table_lines = self._table_rows_to_readable_lines(table_rows)
                current_lines.extend(table_lines)

        if current_lines:
            collected.append(
                {
                    "section_title": current_title,
                    "content": "\n".join(current_lines).strip(),
                }
            )
        return [item for item in collected if item.get("content")]

    def _is_heading_like(self, paragraph: Paragraph, text: str) -> bool:
        style_name = (paragraph.style.name if paragraph.style else "") or ""
        style_lower = style_name.lower()
        if "heading" in style_lower or style_lower in {"title", "subtitle"}:
            return True
        if re.match(r"^\d+(\.\d+)*[\)\.]?\s+[A-Za-z].*", text):
            return True
        if len(text) <= 80 and not text.endswith(".") and text[0].isupper():
            alpha = [c for c in text if c.isalpha()]
            if alpha and (sum(1 for c in alpha if c.isupper()) / len(alpha)) > 0.55:
                return True
        return False

    def _parse_document(
        self,
        document: DocumentType,
        logical_sections: list[dict[str, str]],
    ) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        current: dict[str, Any] | None = None
        preamble: list[str] = []
        recent_context = ""
        last_heading = ""
        pending_key: str | None = None
        pending_table_hint: str | None = None

        def flush_preamble_to(section: dict[str, Any]) -> None:
            nonlocal preamble
            if not preamble:
                return
            blob = "\n".join(preamble).strip()
            preamble = []
            if not blob:
                return
            if section.get("raw_text"):
                section["raw_text"] = section["raw_text"].rstrip() + "\n" + blob + "\n"
            else:
                section["raw_text"] = blob + "\n"

        def append_section(sec: dict[str, Any]) -> None:
            nonlocal current
            sections.append(sec)
            current = sec
            flush_preamble_to(sec)

        def ensure_api_section(api_reference_id: str) -> dict[str, Any]:
            nonlocal current
            for s in sections:
                if s.get("api_reference_id") == api_reference_id:
                    current = s
                    return s
            sec = self._new_section(api_reference_id)
            append_section(sec)
            return sec

        def start_sso_section() -> dict[str, Any]:
            """Synthetic id for OpenID getSSOToken when doc has no REST id on that block."""
            sec = self._new_section("getSSOToken")
            sec["service_name"] = "getSSOToken"
            sec["service_group"] = "Openid"
            append_section(sec)
            return sec

        for block in self._iter_block_items(document):
            if isinstance(block, Paragraph):
                text = self._normalize_text(block.text)
                if not text:
                    continue
                is_heading = self._is_heading_like(block, text)
                if is_heading:
                    last_heading = text
                recent_context = text

                if is_heading and _GET_SSO_TOKEN_HEADING.match(text):
                    if current is None or current.get("api_reference_id") != "getSSOToken":
                        start_sso_section()
                    self._append_raw_text(current, text)
                    pending_key = None
                    continue

                refs = self._extract_api_refs(text)
                if refs:
                    primary = refs[0]
                    ensure_api_section(primary)
                    if current is not None and _QR_CHECKOUT_HEADING.search(text):
                        current["service_name"] = current.get("service_name") or "QR Checkout URL Endpoint"
                        if not current.get("service_group") or "/" in text:
                            current["service_group"] = "/v1/provisioning/generate-qr"
                    pending_key = None

                if current is None:
                    preamble.append(text)
                    pending_key = None
                    continue

                if is_heading and _QR_CHECKOUT_HEADING.search(text) and current.get("api_reference_id"):
                    if API_REF_PATTERN.search(current.get("api_reference_id") or ""):
                        current["service_name"] = current.get("service_name") or "QR Checkout URL Endpoint"
                        if "/" in text or "provisioning" in text.lower():
                            current["service_group"] = "/v1/provisioning/generate-qr"

                if pending_key and self._looks_like_value_line(text):
                    self._map_key_value(current, pending_key, text)
                    pending_key = None
                elif self._looks_like_metadata_key_line(text):
                    pending_key = text
                else:
                    pending_key = None
                self._append_raw_text(current, text)
                self._extract_inline_metadata(current, text)

                cap = text.strip()
                if _RESPONSE_TABLE_CAPTION_RE.match(cap):
                    pending_table_hint = "response"
                elif _REQUEST_TABLE_CAPTION_RE.match(cap):
                    low = cap.lower()
                    if "header" in low:
                        pending_table_hint = "header"
                    elif "query" in low:
                        pending_table_hint = "query"
                    else:
                        pending_table_hint = "request"
                elif _ERROR_TABLE_CAPTION_RE.match(cap):
                    pending_table_hint = "error"
                elif _JWT_TABLE_CAPTION_RE.match(cap):
                    pending_table_hint = "jwt"
            elif isinstance(block, Table):
                table_rows = self._extract_table_rows(block)
                if not table_rows:
                    continue
                table_text = "\n".join(self._table_rows_to_readable_lines(table_rows))
                refs = self._extract_api_refs(table_text)
                if refs:
                    ensure_api_section(refs[0])
                if current is None:
                    preamble.append(table_text)
                    pending_table_hint = None
                    continue
                self._append_raw_text(current, table_text)
                ctx = f"{recent_context} {last_heading}".strip()
                hint = pending_table_hint
                pending_table_hint = None
                self._parse_table_into_section(current, table_rows, ctx, section_hint=hint)

        if sections and preamble:
            flush_preamble_to(sections[0])

        self._apply_logical_section_hints(sections, logical_sections)
        self._apply_service_defaults(sections)
        return sections

    def _apply_service_defaults(self, sections: list[dict[str, Any]]) -> None:
        """Fill typical QR Util / OpenID template fields when the doc omits explicit rows."""
        for section in sections:
            ref = (section.get("api_reference_id") or "").upper()
            name = (section.get("service_name") or "").lower()
            group = (section.get("service_group") or "").lower()
            if ref == "GETSSOTOKEN":
                section.setdefault("service_name", "getSSOToken")
                section.setdefault("service_group", "Openid")
                section.setdefault("service_method", "POST")
                section.setdefault("api_authentication", "client id and secret")
                section.setdefault("api_gateway", "Yes")
            if ref.startswith("API-REST-") and (
                "qr checkout" in name
                or "generate-qr" in group
                or "/v1/provisioning" in (section.get("service_group") or "")
            ):
                section.setdefault("service_name", "QR Checkout URL Endpoint")
                section.setdefault("service_group", "/v1/provisioning/generate-qr")
                section.setdefault("service_method", "POST")
                section.setdefault("service_type", "REST (JSON)")
                section.setdefault("service_pattern", "Synchronous")
                section.setdefault("service_max_timeout", "Default (30 sec)")
                section.setdefault("api_authentication", "Token")
                section.setdefault("api_gateway", "Yes")

    def _apply_logical_section_hints(
        self,
        sections: list[dict[str, Any]],
        logical_sections: list[dict[str, str]],
    ) -> None:
        """Attach OAuth grant narrative to getSSOToken section when present."""
        if not sections:
            return
        sso = next((s for s in sections if s.get("api_reference_id") == "getSSOToken"), None)
        if not sso:
            return
        for sec in logical_sections:
            tl = (sec.get("section_title") or "").lower()
            if "oauth2" in tl and "credential" in tl:
                body = self._normalize_text(sec.get("content") or "")
                if body and body not in (sso.get("raw_text") or ""):
                    self._append_raw_text(sso, f"{sec.get('section_title')}\n{body}")
                if not sso.get("api_authentication") and "client" in body.lower():
                    sso["api_authentication"] = "client id and secret"

    def _extract_document_title(self, document: DocumentType) -> str | None:
        for paragraph in document.paragraphs:
            text = self._normalize_text(paragraph.text)
            if text and len(text) > 5 and not API_REF_PATTERN.search(text):
                return text
        return None

    def _extract_purpose_scope(self, document: DocumentType) -> str | None:
        sections = self._extract_logical_sections(document)
        for section in sections:
            title = (section.get("section_title") or "").lower()
            content = self._normalize_text(section.get("content") or "")
            if not content:
                continue
            if "introduction" in title or "purpose" in title or "scope" in title:
                return content
        for paragraph in document.paragraphs:
            text = self._normalize_text(paragraph.text)
            if not text:
                continue
            lowered = text.lower()
            if "this document details the information that is required to process broadband order requests" in lowered:
                return text
            if "purpose" in lowered or "scope" in lowered:
                return text
        return None

    def _extract_table_rows(self, table: Table) -> list[list[str]]:
        """Preserve empty leading cells so nested parameter rows stay column-aligned."""
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [self._normalize_text(cell.text) for cell in row.cells]
            while len(cells) > 1 and cells[-1] == "":
                cells.pop()
            if any(c != "" for c in cells):
                rows.append(cells)
        return rows

    def _parse_table_into_section(
        self,
        section: dict[str, Any],
        table_rows: list[list[str]],
        context: str,
        *,
        section_hint: str | None = None,
    ) -> None:
        work_rows = list(table_rows)
        hint = section_hint
        if len(work_rows) >= 2 and len(work_rows[0]) == 1:
            title = (work_rows[0][0] or "").strip()
            if _RESPONSE_TABLE_CAPTION_RE.match(title):
                hint = hint or "response"
                work_rows = work_rows[1:]
            elif _REQUEST_TABLE_CAPTION_RE.match(title):
                tl = title.lower()
                if "header" in tl:
                    hint = hint or "header"
                elif "query" in tl:
                    hint = hint or "query"
                else:
                    hint = hint or "request"
                work_rows = work_rows[1:]
            elif _ERROR_TABLE_CAPTION_RE.match(title):
                hint = hint or "error"
                work_rows = work_rows[1:]
            elif _JWT_TABLE_CAPTION_RE.match(title):
                hint = hint or "jwt"
                work_rows = work_rows[1:]

        for key, value in self._table_rows_to_key_values(work_rows):
            self._map_key_value(section, key, value)

        max_cols = max((len(row) for row in work_rows), default=0)
        if max_cols <= 2:
            return

        header_row = [self._normalize_key(col) for col in work_rows[0]]
        joined_headers = " ".join(header_row)
        is_param_table = any(k in joined_headers for k in ("param", "field", "element", "attribute")) or (
            "name" in joined_headers and any(x in joined_headers for x in ("mandatory", "type", "description", "optional"))
        ) or ("field" in joined_headers and "type" in joined_headers)

        context_lower = context.lower()

        def infer_route() -> str | None:
            if hint:
                return hint
            if "jwt" in context_lower or "payload structure" in context_lower or "token payload" in context_lower:
                return "jwt"
            if "error" in context_lower and ("code" in context_lower or "expected" in context_lower):
                return "error"
            if "header" in context_lower:
                return "header"
            if "query" in context_lower:
                return "query"
            if "output" in context_lower or "response" in context_lower:
                return "response"
            if "request" in context_lower or "body" in context_lower:
                return "request"
            return None

        request_parent = ""
        nested_response_prefix = ""

        def append_by_route(route: str | None, entry: dict[str, Any]) -> None:
            r = route or "request"
            if r == "jwt":
                section.setdefault("jwt_payload_parameters", []).append(entry)
            elif r == "error":
                section.setdefault("error_code_parameters", []).append(entry)
            elif r == "header":
                section["header_parameters"].append(entry)
            elif r == "query":
                section.setdefault("query_parameters", []).append(entry)
            elif r == "response":
                section["output_response_success"].append(entry)
                section["api_response_parameters"].append(entry)
            else:
                section["input_parameters"].append(entry)
                section["api_request_parameters"].append(entry)

        if not is_param_table:
            for row in work_rows[1:]:
                if len(row) < 2:
                    continue
                while len(row) < 5:
                    row.append("")
                section["input_parameters"].append(
                    {
                        "param_name": (row[0] or "").strip() or None,
                        "param_type": (row[1] or "").strip() or None,
                        "mandatory_optional": (row[2] or "").strip() or None,
                        "description": (row[3] or "").strip() or None,
                    }
                )
            return

        is_response_ctx = hint == "response" or "output" in context_lower or "response" in context_lower
        is_request_ctx = hint in {"request", "header", "query"} or (
            "request" in context_lower or "header" in context_lower or "query" in context_lower or "body" in context_lower
        )

        for row in work_rows[1:]:
            if not any((c or "").strip() for c in row):
                continue
            while len(row) < 5:
                row.append("")
            name_raw = (row[0] or "").strip()
            c1, c2, c3, c4 = (row[1] or "").strip(), (row[2] or "").strip(), (row[3] or "").strip(), (row[4] or "").strip()

            def build_entry(pname: str, ptype: str, mand: str, desc: str) -> dict[str, Any]:
                return {
                    "param_name": pname or None,
                    "param_type": ptype or None,
                    "mandatory_optional": mand or None,
                    "description": desc or None,
                }

            type_guess = (c1 or "").strip().lower()
            object_like = type_guess in ("object", "array", "json", "{}")

            pname = name_raw
            ptype, mand, desc = c1, c2, c3

            if is_response_ctx:
                if name_raw and object_like:
                    nested_response_prefix = name_raw
                    pname = name_raw
                elif nested_response_prefix and name_raw:
                    pname = f"{nested_response_prefix}.{name_raw}"
                elif not name_raw:
                    inner = c1
                    if inner and nested_response_prefix:
                        pname = f"{nested_response_prefix}.{inner}"
                        ptype, mand, desc = c2, c3, c4
                    elif inner:
                        pname = inner
                        ptype, mand, desc = c2, c3, c4
                else:
                    pname = name_raw
            elif is_request_ctx and not name_raw and request_parent:
                inner = c1
                if inner:
                    pname = f"{request_parent}.{inner}"
                    ptype, mand, desc = c2, c3, c4
            elif is_request_ctx and name_raw:
                pname = name_raw
                ptype, mand, desc = c1, c2, c3
                if object_like:
                    request_parent = name_raw
                else:
                    request_parent = ""

            entry = build_entry(pname, ptype, mand, desc)
            route = infer_route()
            append_by_route(route, entry)

    def _extract_inline_metadata(self, section: dict[str, Any], text: str) -> None:
        if ":" not in text:
            return
        key, value = text.split(":", 1)
        self._map_key_value(section, key, value)

    def _new_section(self, api_reference_id: str) -> dict[str, Any]:
        return {
            "api_reference_id": api_reference_id,
            "service_name": None,
            "service_group": None,
            "service_description": None,
            "service_swagger": None,
            "service_url": None,
            "api_authentication": None,
            "api_gateway": None,
            "source": None,
            "service_method": None,
            "service_type": None,
            "service_pattern": None,
            "service_max_timeout": None,
            "header_parameters": [],
            "input_parameters": [],
            "query_parameters": [],
            "jwt_payload_parameters": [],
            "error_code_parameters": [],
            "output_response_success": [],
            "api_request_parameters": [],
            "api_response_parameters": [],
            "sample_initial": None,
            "sample_token_exist": None,
            "sample_request": None,
            "sample_success_response": None,
            "sample_failed_request": None,
            "sample_failed_response": None,
            "raw_text": "",
        }

    def _map_key_value(self, section: dict[str, Any], key: str, value: str) -> None:
        key_norm = self._normalize_key(key)
        value_norm = self._normalize_text(value)
        if not value_norm:
            return
        if value_norm.lower() == key_norm:
            return
        mapping = {
            "service name": "service_name",
            "service group": "service_group",
            "service description": "service_description",
            "service swagger": "service_swagger",
            "service swagger url": "service_swagger",
            "service url": "service_url",
            "api authentication": "api_authentication",
            "authentication": "api_authentication",
            "api gateway": "api_gateway",
            "source": "source",
            "service method": "service_method",
            "service type": "service_type",
            "service pattern": "service_pattern",
            "service max timeout": "service_max_timeout",
            "max timeout": "service_max_timeout",
            "sample initial": "sample_initial",
            "sample token exist": "sample_token_exist",
            "sample request": "sample_request",
            "sample success response": "sample_success_response",
            "sample failed request": "sample_failed_request",
            "sample failed response": "sample_failed_response",
        }
        target = mapping.get(key_norm)
        if target and value_norm.lower() not in mapping:
            section[target] = value_norm

    def _extract_api_refs(self, text: str) -> list[str]:
        return [match.upper() for match in API_REF_PATTERN.findall(text or "")]

    def _table_rows_to_key_values(self, table_rows: list[list[str]]) -> list[tuple[str, str]]:
        pairs: list[tuple[str, str]] = []
        recognized_keys = {
            "service name",
            "service group",
            "service description",
            "service method",
            "service type",
            "service pattern",
            "service max timeout",
            "api authentication",
            "api gateway",
            "source",
            "sample request",
            "sample success response",
            "sample failed request",
            "sample failed response",
        }
        consumed_rows: set[int] = set()
        for idx, row in enumerate(table_rows):
            if idx in consumed_rows:
                continue
            if len(row) >= 2:
                key = row[0]
                key_norm = self._normalize_key(key)
                value_candidate = next(
                    (
                        cell
                        for cell in row[1:]
                        if self._normalize_key(cell)
                        and self._normalize_key(cell) != key_norm
                        and self._normalize_key(cell) not in recognized_keys
                    ),
                    None,
                )
                value = value_candidate or row[1]
                value_norm = self._normalize_key(value)
                if key_norm in recognized_keys and (not value_norm or value_norm in recognized_keys or value_norm == key_norm):
                    if idx + 1 < len(table_rows):
                        next_row = table_rows[idx + 1]
                        next_value = self._normalize_text(" | ".join(next_row))
                        next_value_norm = self._normalize_key(next_value)
                        if next_value and next_value_norm not in recognized_keys:
                            pairs.append((key, next_value))
                            consumed_rows.add(idx + 1)
                    continue
                pairs.append((key, value))

        linear_cells = [row[0] for index, row in enumerate(table_rows) if len(row) == 1 and index not in consumed_rows]
        pending_key: str | None = None
        for cell in linear_cells:
            cell_norm = self._normalize_key(cell)
            if pending_key is None and cell_norm in recognized_keys:
                pending_key = cell
                continue
            if pending_key and cell_norm not in recognized_keys:
                pairs.append((pending_key, cell))
                pending_key = None
        return pairs

    def _table_rows_to_readable_lines(self, table_rows: list[list[str]]) -> list[str]:
        lines: list[str] = []
        for row in table_rows:
            if len(row) >= 2:
                lines.append(f"{row[0]}: {' | '.join(row[1:])}")
            elif len(row) == 1:
                lines.append(row[0])
        return lines

    def _looks_like_metadata_key_line(self, text: str) -> bool:
        key = self._normalize_key(text)
        return key in {
            "service name",
            "service group",
            "service description",
            "service method",
            "service type",
            "service pattern",
            "service max timeout",
            "api authentication",
            "api gateway",
            "source",
            "sample request",
            "sample success response",
            "sample failed request",
            "sample failed response",
        }

    def _looks_like_value_line(self, text: str) -> bool:
        return bool(text and ":" not in text and len(text) > 2)

    def _append_raw_text(self, section: dict[str, Any], text: str) -> None:
        if not text:
            return
        if section["raw_text"]:
            section["raw_text"] += "\n"
        section["raw_text"] += text

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text or "").strip()

    def _normalize_key(self, key: str) -> str:
        return self._normalize_text(key).lower()

    def _iter_block_items(self, document: DocumentType) -> Iterator[Paragraph | Table]:
        parent_elm = document.element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)


docx_parser_service = DocxParserService()
