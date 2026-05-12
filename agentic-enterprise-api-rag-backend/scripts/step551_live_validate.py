#!/usr/bin/env python3
"""Step 55.1 — Re-ingest QR-style DOCX (fixture-based + enrichment lines) and live-validate."""

from __future__ import annotations

import json
import os
import sys
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document
from dotenv import load_dotenv

BACKEND_ROOT = Path(__file__).resolve().parent.parent


def _build_step551_docx_bytes() -> bytes:
    """QR Util-style DOCX: based on test fixture + explicit expiry + second header for validation."""
    doc = Document()
    doc.add_heading("QR Util Service API", level=1)
    doc.add_paragraph("QR provisioning and checkout.")
    doc.add_heading("General Authentication", level=1)
    doc.add_paragraph("The QR Util Service uses OAuth2 Client Credentials Grant for service-to-service calls.")
    doc.add_paragraph(
        "Non-prod access tokens expire in 540 seconds. Prod access tokens expire in 60 seconds."
    )
    doc.add_heading("OAuth2 Client Credentials Grant", level=1)
    doc.add_paragraph("Obtain an access token using a registered client id and secret.")
    doc.add_heading("getSSOToken", level=1)
    sso = doc.add_table(rows=4, cols=2)
    for idx, (k, v) in enumerate(
        [
            ("Service Method", "POST"),
            ("Service Group", "Openid"),
            ("API Authentication", "client id and secret"),
            ("API Gateway", "Yes"),
        ]
    ):
        sso.rows[idx].cells[0].text = k
        sso.rows[idx].cells[1].text = v
    doc.add_heading("API-REST-QRU-01", level=1)
    doc.add_heading("QR Checkout URL Endpoint", level=1)
    qr = doc.add_table(rows=8, cols=2)
    for idx, (k, v) in enumerate(
        [
            ("Service Name", "QR Checkout URL Endpoint"),
            ("Service Group", "/v1/provisioning/generate-qr"),
            ("Service Method", "POST"),
            ("Service Type", "REST (JSON)"),
            ("Service Pattern", "Synchronous"),
            ("Service Max Timeout", "Default (30 sec)"),
            ("API Authentication", "Token / Bearer token"),
            ("API Gateway", "Yes"),
        ]
    ):
        qr.rows[idx].cells[0].text = k
        qr.rows[idx].cells[1].text = v

    def add_param_block(title: str, p1: str, p2: str) -> None:
        doc.add_paragraph(title)
        t = doc.add_table(rows=2, cols=4)
        for c, name in enumerate(["Name", "Type", "Mandatory", "Description"]):
            t.rows[0].cells[c].text = name
        t.rows[1].cells[0].text = p1
        t.rows[1].cells[1].text = "string"
        t.rows[1].cells[2].text = "Mandatory"
        t.rows[1].cells[3].text = p2

    add_param_block("Header Parameters", "Authorization", "bearer + token from IDP")
    add_param_block("Header Parameters", "TransactionId", "API-<method>-<DDMMYYYY>-<UUID>")
    add_param_block("Query Parameters", "source", "client application LOV")
    doc.add_paragraph("Input Parameter")
    doc.add_paragraph("Request Parameters")
    it = doc.add_table(rows=5, cols=4)
    for c, name in enumerate(["Name", "Type", "Mandatory", "Description"]):
        it.rows[0].cells[c].text = name
    req_rows = [
        ("source", "String", "Mandatory", "client application"),
        ("targetUrl", "String", "Mandatory", "destination"),
        ("request_type", "String", "Mandatory", "operation"),
        ("parameters", "Array", "Mandatory", "key/value array"),
    ]
    for i, (a, b, c2, d) in enumerate(req_rows, start=1):
        it.rows[i].cells[0].text = a
        it.rows[i].cells[1].text = b
        it.rows[i].cells[2].text = c2
        it.rows[i].cells[3].text = d

    doc.add_paragraph("Response Parameters")
    rt = doc.add_table(rows=12, cols=4)
    for c, name in enumerate(["Name", "Type", "Mandatory", "Description"]):
        rt.rows[0].cells[c].text = name
    resp_rows = [
        ("status", "string", "Mandatory", "Overall status"),
        ("code", "string", "Mandatory", "Status code"),
        ("desc", "string", "Mandatory", "Description"),
        ("timestamp", "string", "Mandatory", "Timestamp"),
        ("transactionId", "string", "Mandatory", "Transaction id"),
        ("correlationId", "string", "Mandatory", "Correlation id"),
        ("responseInfo", "Object", "Conditional", "Nested response payload"),
        ("", "merchantName", "string", "Optional"),
        ("qrText", "string", "Mandatory", "Encoded QR payload"),
        ("errorcode", "string", "Optional", "Error code when failed"),
        ("errormsg", "string", "Optional", "Error message when failed"),
    ]
    for i, (a, b, c2, d) in enumerate(resp_rows, start=1):
        rt.rows[i].cells[0].text = a
        rt.rows[i].cells[1].text = b
        rt.rows[i].cells[2].text = c2
        rt.rows[i].cells[3].text = d

    add_param_block("Expected Error Codes", "400", "Bad request when mandatory fields missing")
    add_param_block("JWT payload structure", "sub", "Subject identifier")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def main() -> int:
    load_dotenv(BACKEND_ROOT / ".env")
    base = os.getenv("STEP551_API_BASE_URL", "http://127.0.0.1:8010/api/v1").rstrip("/")
    kb_id = int(os.getenv("STEP551_KB_ID") or os.getenv("CLEANUP_KB_ID") or "1")
    timeout = float(os.getenv("STEP551_TIMEOUT", "180"))

    jwt = (os.getenv("DEMO_SMOKE_JWT") or os.getenv("CLEANUP_JWT") or "").strip()
    email = (os.getenv("DEMO_SMOKE_EMAIL") or "superadmin@local").strip()
    password = (os.getenv("DEMO_SMOKE_PASSWORD") or os.getenv("CLEANUP_PASSWORD") or "").strip()

    headers: dict[str, str] = {}
    if jwt:
        headers["Authorization"] = f"Bearer {jwt}"
    elif password:
        with httpx.Client(timeout=30.0) as client:
            r = client.post(f"{base}/auth/login", json={"email": email, "password": password})
        if r.status_code != 200:
            print(json.dumps({"error": "login_failed", "status": r.status_code, "body": r.text[:300]}, indent=2))
            return 1
        tok = (r.json() or {}).get("access_token")
        if not tok:
            print(json.dumps({"error": "no_token"}, indent=2))
            return 1
        headers["Authorization"] = f"Bearer {tok}"
    else:
        print("ERROR: Set DEMO_SMOKE_PASSWORD or DEMO_SMOKE_JWT", file=sys.stderr)
        return 1

    out: dict[str, object] = {"kb_id": kb_id}

    # Resolve project_id from existing KB documents (safest before cleanup).
    project_id = int(os.getenv("STEP551_PROJECT_ID") or "0")
    with httpx.Client(timeout=timeout) as client:
        dr = client.get(f"{base}/ingestion/documents", params={"knowledge_base_id": kb_id}, headers=headers)
        if dr.status_code == 200:
            docs = (dr.json() or {}).get("documents") or []
            if docs:
                project_id = int(docs[0].get("project_id") or 0)
    if project_id <= 0:
        project_id = 1
    out["project_id_used"] = project_id

    # Cleanup (destructive for this KB only).
    cleanup_body = {
        "knowledge_base_id": kb_id,
        "delete_documents": True,
        "delete_vectors": True,
        "delete_chat_sessions": True,
        "delete_feedback": False,
        "delete_improvement_tasks": False,
        "include_demo_seed": False,
        "delete_audit_logs": False,
        "dry_run": False,
    }
    with httpx.Client(timeout=180.0) as client:
        cr = client.post(
            f"{base}/admin/test-data/cleanup",
            headers={**headers, "Content-Type": "application/json"},
            json=cleanup_body,
        )
        try:
            out["cleanup"] = cr.json()
        except Exception:
            out["cleanup"] = {"status_code": cr.status_code, "text": cr.text[:500]}
        if cr.status_code != 200:
            print(json.dumps(out, indent=2, default=str))
            return 1

        # Ingest
        file_bytes = _build_step551_docx_bytes()
        files = {"file": ("step551_qr_util.docx", file_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {
            "project_id": str(project_id),
            "knowledge_base_id": str(kb_id),
            "document_type": "api",
        }
        ir = client.post(f"{base}/ingestion/ingest-docx", headers=headers, files=files, data=data)
        try:
            out["ingestion"] = ir.json()
        except Exception:
            out["ingestion"] = {"status_code": ir.status_code, "text": ir.text[:800]}
        if ir.status_code != 200:
            print(json.dumps(out, indent=2, default=str))
            return 1

        # Chunk type breakdown via diagnostics SQL-free: retrieval-test exposes chunk types in chunks[].chunk_type
        doc_id = (ir.json() or {}).get("document_id")

        # KB id for queries
        kr = client.get(f"{base}/knowledge-bases/me", headers=headers)
        klist = (kr.json() if kr.status_code == 200 else []) or []
        kb_ids = [int(x["id"]) for x in klist if int(x.get("document_count") or 0) > 0]
        ask_kb = kb_id if kb_id in kb_ids else (kb_ids[0] if kb_ids else kb_id)

        def retrieval_test(question: str) -> dict[str, object]:
            rr = client.post(
                f"{base}/diagnostics/retrieval-test",
                headers=headers,
                json={"knowledge_base_id": ask_kb, "question": question, "top_k": 8},
            )
            try:
                return rr.json() if rr.content else {"http_error": rr.status_code, "text": rr.text[:400]}
            except Exception:
                return {"parse_error": rr.text[:400]}

        def ask(question: str) -> dict[str, object]:
            ar = client.post(
                f"{base}/query/ask",
                headers=headers,
                json={
                    "project_id": project_id,
                    "knowledge_base_id": ask_kb,
                    "question": question,
                    "top_k": 8,
                    "debug": True,
                },
                timeout=timeout,
            )
            try:
                return ar.json() if ar.content else {"http_error": ar.status_code, "text": ar.text[:400]}
            except Exception:
                return {"parse_error": ar.text[:400]}

        q1 = "What headers are required while calling the QR generation API?"
        q2 = "What is the expiry time for access tokens in PROD vs non-PROD?"
        q3 = "Explain the request structure for QR generation."

        r1 = retrieval_test(q1)
        r2 = retrieval_test(q2)
        r3 = retrieval_test(q3)

        def chunk_summary(ret: dict[str, object]) -> dict[str, object]:
            chunks = ret.get("chunks") if isinstance(ret, dict) else None
            types: list[str] = []
            if isinstance(chunks, list):
                for c in chunks:
                    if isinstance(c, dict) and c.get("chunk_type"):
                        types.append(str(c["chunk_type"]))
            diag = ret.get("diagnostics") if isinstance(ret, dict) else {}
            intents = []
            if isinstance(diag, dict):
                intents = diag.get("detected_intents") or []
            return {"chunk_types_top": types[:8], "detected_intents": intents}

        out["retrieval_Q1"] = {"question": q1, **chunk_summary(r1), "raw_ok": isinstance(r1, dict)}
        out["retrieval_Q2"] = {"question": q2, **chunk_summary(r2), "raw_ok": isinstance(r2, dict)}
        out["retrieval_Q3"] = {"question": q3, **chunk_summary(r3), "raw_ok": isinstance(r3, dict)}

        a1 = ask(q1)
        a2 = ask(q2)
        a3 = ask(q3)

        def ask_compact(a: dict[str, object]) -> dict[str, object]:
            diag = a.get("diagnostics") if isinstance(a, dict) else {}
            recovered = None
            intents = []
            if isinstance(diag, dict):
                recovered = diag.get("response_fields_recovered_from_generic")
                intents = diag.get("detected_intents") or diag.get("prompt_context_diagnostics", {}).get("detected_intents") or []
                if not intents:
                    vr = diag.get("vector_retrieval_diagnostics") or {}
                    if isinstance(vr, dict):
                        intents = vr.get("detected_intents") or []
            return {
                "llm_status": a.get("llm_status") if isinstance(a, dict) else None,
                "answer_preview": ((a.get("answer") or "")[:900] if isinstance(a, dict) else ""),
                "response_fields_recovered_from_generic": recovered,
                "detected_intents_guess": intents,
            }

        out["ask_Q1"] = ask_compact(a1)
        out["ask_Q2"] = ask_compact(a2)
        out["ask_Q3"] = ask_compact(a3)

        reg = [
            "What authentication is required for the QR generation API?",
            "What are the error codes for the QR generation API?",
            "What are the success response fields for the QR generation API?",
        ]
        out["regression"] = []
        for rq in reg:
            ax = ask(rq)
            d = ax.get("diagnostics") if isinstance(ax, dict) else {}
            rec = None
            if isinstance(d, dict):
                rec = d.get("response_fields_recovered_from_generic")
            out["regression"].append(
                {
                    "question": rq,
                    "llm_status": ax.get("llm_status") if isinstance(ax, dict) else None,
                    "response_fields_recovered_from_generic": rec,
                }
            )

        out["document_id"] = doc_id

    print(json.dumps(out, indent=2, default=str))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
