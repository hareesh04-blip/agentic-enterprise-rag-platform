from io import BytesIO

from docx import Document

from app.services.chunking_service import ApiChunkingService, create_document_chunks, response_parameters_chunk_quality
from app.services.docx_parser_service import docx_parser_service
from app.services.query_intent_service import detect_query_intents
from app.services.rag_service import RagService
from app.services.retrieval_service import RetrievalService
from app.services.suggested_question_service import suggested_question_service


def _build_sample_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This document details the information that is required to Process Broadband Order requests from ISOTON to DSE through Application."
    )
    doc.add_heading("API-REST-DSE-01", level=1)
    table = doc.add_table(rows=9, cols=2)
    rows = [
        ("Service Name", "getAppointment"),
        ("Service Group", "broadband"),
        ("Service Description", "To get the list of valid appointment slots when NLT appointment is required."),
        ("Service Method", "POST"),
        ("Service Pattern", "/appointments/get"),
        ("API Authentication", "Bearer"),
        ("API Gateway", "DSE"),
        ("Source", "BB Order Service"),
        ("Sample Request", '{"customerId":"123"}'),
    ]
    for idx, (k, v) in enumerate(rows):
        table.rows[idx].cells[0].text = k
        table.rows[idx].cells[1].text = v

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def test_docx_extraction_includes_table_metadata() -> None:
    parsed = docx_parser_service.parse_preview("sample.docx", _build_sample_docx_bytes())
    api = next((item for item in parsed["apis_full"] if item.get("api_reference_id") == "API-REST-DSE-01"), None)
    assert api is not None
    assert api.get("service_name") == "getAppointment"
    assert "appointment slots" in (api.get("service_description") or "").lower()


def test_purpose_of_document_chunk_created() -> None:
    parsed = docx_parser_service.parse_preview("sample.docx", _build_sample_docx_bytes())
    chunks = create_document_chunks(parsed, document_type="api")
    overview = next((chunk for chunk in chunks if chunk.get("chunk_type") == "document_overview_chunk"), None)
    assert overview is not None
    assert "This document details the information that is required to Process Broadband Order requests" in (
        overview.get("chunk_text") or ""
    )


def test_api_reference_chunk_contains_service_name() -> None:
    parsed = docx_parser_service.parse_preview("sample.docx", _build_sample_docx_bytes())
    chunks = create_document_chunks(parsed, document_type="api")
    metadata_chunk = next(
        (
            chunk
            for chunk in chunks
            if chunk.get("chunk_type") == "api_metadata_chunk"
            and "API Reference ID: API-REST-DSE-01" in (chunk.get("chunk_text") or "")
        ),
        None,
    )
    assert metadata_chunk is not None
    assert "Service Name: getAppointment" in (metadata_chunk.get("chunk_text") or "")


def _build_qr_util_style_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("QR Util Service API", level=1)
    doc.add_paragraph("QR provisioning and checkout.")
    doc.add_heading("General Authentication", level=1)
    doc.add_paragraph("The QR Util Service uses OAuth2 Client Credentials Grant for service-to-service calls.")
    doc.add_heading("OAuth2 Client Credentials Grant", level=1)
    doc.add_paragraph("Obtain an access token using a registered client id and secret.")
    doc.add_heading("getSSOToken", level=1)
    sso = doc.add_table(rows=4, cols=2)
    for idx, (k, v) in enumerate(
        [
            ("Service Method", "POST"),
            ("Service Group", "Openid"),
            ("API Authentication", "client id and secret"),
            ("API Gateway", "Yes"),
        ]
    ):
        sso.rows[idx].cells[0].text = k
        sso.rows[idx].cells[1].text = v
    doc.add_heading("API-REST-QRU-01", level=1)
    doc.add_heading("QR Checkout URL Endpoint", level=1)
    qr = doc.add_table(rows=8, cols=2)
    for idx, (k, v) in enumerate(
        [
            ("Service Name", "QR Checkout URL Endpoint"),
            ("Service Group", "/v1/provisioning/generate-qr"),
            ("Service Method", "POST"),
            ("Service Type", "REST (JSON)"),
            ("Service Pattern", "Synchronous"),
            ("Service Max Timeout", "Default (30 sec)"),
            ("API Authentication", "Token"),
            ("API Gateway", "Yes"),
        ]
    ):
        qr.rows[idx].cells[0].text = k
        qr.rows[idx].cells[1].text = v

    def add_param_block(title: str, p1: str, p2: str) -> None:
        doc.add_paragraph(title)
        t = doc.add_table(rows=2, cols=4)
        for c, name in enumerate(["Name", "Type", "Mandatory", "Description"]):
            t.rows[0].cells[c].text = name
        t.rows[1].cells[0].text = p1
        t.rows[1].cells[1].text = "string"
        t.rows[1].cells[2].text = "Mandatory"
        t.rows[1].cells[3].text = p2

    add_param_block("Header Parameters", "Authorization", "Bearer access token")
    add_param_block("Query Parameters", "merchantId", "Merchant identifier")

    doc.add_paragraph("Response Parameters")
    rt = doc.add_table(rows=12, cols=4)
    for c, name in enumerate(["Name", "Type", "Mandatory", "Description"]):
        rt.rows[0].cells[c].text = name
    resp_rows = [
        ("status", "string", "Mandatory", "Overall status"),
        ("code", "string", "Mandatory", "Status code"),
        ("desc", "string", "Mandatory", "Description"),
        ("timestamp", "string", "Mandatory", "Timestamp"),
        ("transactionId", "string", "Mandatory", "Transaction id"),
        ("correlationId", "string", "Mandatory", "Correlation id"),
        ("responseInfo", "Object", "Conditional", "Nested response payload"),
        ("", "merchantName", "string", "Optional"),
        ("qrText", "string", "Mandatory", "Encoded QR payload"),
        ("errorcode", "string", "Optional", "Error code when failed"),
        ("errormsg", "string", "Optional", "Error message when failed"),
    ]
    for i, (a, b, c2, d) in enumerate(resp_rows, start=1):
        rt.rows[i].cells[0].text = a
        rt.rows[i].cells[1].text = b
        rt.rows[i].cells[2].text = c2
        rt.rows[i].cells[3].text = d

    add_param_block("Expected Error Codes", "400", "Bad request when mandatory fields missing")
    add_param_block("JWT payload structure", "sub", "Subject identifier")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def test_qr_util_style_docx_parsing_and_chunking() -> None:
    parsed = docx_parser_service.parse_preview("qr_util.docx", _build_qr_util_style_docx_bytes())
    assert parsed["api_count"] >= 1
    assert "OAuth2" in (parsed.get("authentication_preamble") or "")
    apis = {a["api_reference_id"]: a for a in parsed["apis_full"]}
    assert "API-REST-QRU-01" in apis
    assert apis["API-REST-QRU-01"].get("service_name") == "QR Checkout URL Endpoint"
    assert apis["API-REST-QRU-01"].get("service_group") == "/v1/provisioning/generate-qr"
    assert "getSSOToken" in apis
    assert apis["getSSOToken"].get("service_group") == "Openid"
    assert apis["getSSOToken"].get("service_method") == "POST"
    qr = apis["API-REST-QRU-01"]
    assert qr.get("header_parameters")
    assert qr.get("query_parameters")
    assert qr.get("api_response_parameters") or qr.get("output_response_success")
    assert qr.get("error_code_parameters")
    assert qr.get("jwt_payload_parameters")

    chunks = create_document_chunks(parsed, document_type="api")
    assert len(chunks) >= 12
    auth = next((c for c in chunks if c.get("chunk_type") == "authentication_chunk"), None)
    assert auth is not None
    assert "General Authentication" in (auth.get("chunk_text") or "")
    assert "OAuth2" in (auth.get("chunk_text") or "")
    assert any(c.get("chunk_type") == "api_query_parameters_chunk" for c in chunks)
    assert any(c.get("chunk_type") == "api_error_codes_chunk" for c in chunks)
    assert any(c.get("chunk_type") == "api_jwt_payload_chunk" for c in chunks)


def test_step_50_6_fallback_scoped_past_toc_no_auth_noise() -> None:
    svc = ApiChunkingService()
    raw = (
        "Table of Contents\n"
        "Response Parameters ........................................... 9\n\n"
        "API-REST-QRU-01\n"
        "QR Checkout URL Endpoint\n"
        "General Authentication\n"
        "Authentication without JWT Signing — do not capture this.\n\n"
        "Response Parameters\n"
        "Name: Type: Mandatory: Description\n"
        "status: string: Mandatory: S\n"
        "transactionId: string: Mandatory: T\n"
        "correlationId: string: Mandatory: C\n"
        "responseInfo: Object: Conditional: O\n"
        "Expected Error Codes\n"
        "400\n"
    )
    fb = svc._fallback_response_params_from_raw(raw, "API-REST-QRU-01")
    assert fb is not None
    assert "Table of Contents" not in fb
    assert "Authentication without JWT" not in fb
    assert "transactionId" in fb
    assert "correlationId" in fb
    assert "responseInfo" in fb


def test_step_50_6_fallback_rejected_without_two_known_fields() -> None:
    svc = ApiChunkingService()
    raw = (
        "API-REST-QRU-01\n"
        "Response Parameters\n"
        "onlyunknownfield: x: y: z\n"
        "Expected Error Codes\n"
    )
    assert svc._fallback_response_params_from_raw(raw, "API-REST-QRU-01") is None


def test_response_parameters_chunk_quality_flags_degenerate_heading() -> None:
    low = response_parameters_chunk_quality(
        "Section: API Response Parameters\nResponse Parameters: response parameters\nGeneral Authentication\n"
    )
    assert not low["boost_ok"]


def test_step_50_6_docx_toc_does_not_pollute_response_chunk() -> None:
    doc = Document()
    doc.add_heading("QR Util", level=1)
    doc.add_paragraph("Table of Contents")
    doc.add_paragraph("Response Parameters ........................................... 12")
    doc.add_heading("General Authentication", level=1)
    doc.add_paragraph("Authentication without JWT Signing placeholder.")
    doc.add_heading("API-REST-QRU-01", level=1)
    doc.add_heading("QR Checkout URL Endpoint", level=1)
    mt = doc.add_table(rows=4, cols=2)
    for idx, (k, v) in enumerate(
        [
            ("Service Name", "QR Checkout URL Endpoint"),
            ("Service Group", "/v1/provisioning/generate-qr"),
            ("Service Method", "POST"),
            ("API Gateway", "Yes"),
        ]
    ):
        mt.rows[idx].cells[0].text = k
        mt.rows[idx].cells[1].text = v
    doc.add_paragraph("Response Parameters")
    rt = doc.add_table(rows=6, cols=4)
    for c, name in enumerate(["Name", "Type", "Mandatory", "Description"]):
        rt.rows[0].cells[c].text = name
    rows = [
        ("status", "string", "Mandatory", "S"),
        ("transactionId", "string", "Mandatory", "T"),
        ("correlationId", "string", "Mandatory", "C"),
        ("responseInfo", "Object", "Conditional", "O"),
        ("qrText", "string", "Mandatory", "Q"),
    ]
    for i, (a, b, c2, d) in enumerate(rows, start=1):
        rt.rows[i].cells[0].text = a
        rt.rows[i].cells[1].text = b
        rt.rows[i].cells[2].text = c2
        rt.rows[i].cells[3].text = d
    bio = BytesIO()
    doc.save(bio)
    parsed = docx_parser_service.parse_preview("toc_qr.docx", bio.getvalue())
    chunks = create_document_chunks(parsed, document_type="api")
    resp = next(c for c in chunks if c.get("chunk_type") == "api_response_parameters_chunk" and "API-REST-QRU-01" in (c.get("chunk_text") or ""))
    txt = resp.get("chunk_text") or ""
    assert "Authentication without JWT Signing" not in txt
    assert "transactionId" in txt
    assert "correlationId" in txt
    assert "responseInfo" in txt
    assert "qrText" in txt


def test_response_field_intent_detected_on_qr_generation_question() -> None:
    q = "What are the success response fields for the QR generation API?"
    intents = detect_query_intents(q)
    assert "response_field_intent" in intents


def test_step_50_5_qr_response_table_extracted_not_na_chunk() -> None:
    """Structured Response Parameters rows populate api_response_parameters (not only generic_section)."""
    parsed = docx_parser_service.parse_preview("qr_step505.docx", _build_qr_util_style_docx_bytes())
    qr = next(x for x in parsed["apis_full"] if x.get("api_reference_id") == "API-REST-QRU-01")
    assert qr.get("api_response_parameters"), "expected parsed api_response_parameters list"
    chunks = create_document_chunks(parsed, document_type="api")
    resp = next(c for c in chunks if c.get("chunk_type") == "api_response_parameters_chunk" and "API-REST-QRU-01" in (c.get("chunk_text") or ""))
    txt = resp.get("chunk_text") or ""
    assert "Response Parameters: N/A" not in txt
    for needle in ("transactionId", "correlationId", "responseInfo", "qrText"):
        assert needle in txt


def test_qr_response_parameters_fields_in_chunk() -> None:
    parsed = docx_parser_service.parse_preview("qr.docx", _build_qr_util_style_docx_bytes())
    qr = next(x for x in parsed["apis_full"] if x.get("api_reference_id") == "API-REST-QRU-01")
    names = " ".join((p.get("param_name") or "") for p in (qr.get("api_response_parameters") or []))
    for field in (
        "status",
        "code",
        "desc",
        "timestamp",
        "transactionId",
        "correlationId",
        "responseInfo",
        "responseInfo.merchantName",
        "qrText",
        "errorcode",
        "errormsg",
    ):
        assert field in names
    chunks = create_document_chunks(parsed, document_type="api")
    resp_chunks = [c for c in chunks if c.get("chunk_type") == "api_response_parameters_chunk"]
    assert resp_chunks
    combined = " ".join((c.get("chunk_text") or "") for c in resp_chunks)
    for field in ("status", "code", "desc", "timestamp", "transactionId", "correlationId", "responseInfo", "qrText", "errorcode", "errormsg"):
        assert field in combined


def test_rerank_prefers_response_chunk_for_response_field_intent() -> None:
    svc = RetrievalService()
    q = "What are the success response fields for the QR generation API?"
    intents = detect_query_intents(q)
    candidates = [
        {
            "chunk_type": "generic_section_chunk",
            "chunk_text": "General notes about provisioning.",
            "vector_score_raw": 0.55,
            "fallback_score_raw": 0.0,
            "score": 0.55,
        },
        {
            "chunk_type": "api_response_parameters_chunk",
            "chunk_text": (
                "API Reference ID: API-REST-QRU-01\nService Name: QR Checkout URL Endpoint\n"
                "Section: API Response Parameters\n"
                "Response Parameters: status (string, Mandatory): ok; code (string, Mandatory): x"
            ),
            "vector_score_raw": 0.22,
            "fallback_score_raw": 0.0,
            "score": 0.22,
            "service_name": "QR Checkout URL Endpoint",
        },
    ]
    ranked = svc._rerank_candidates(candidates=candidates, question=q, intents=intents)
    assert ranked[0]["chunk_type"] == "api_response_parameters_chunk"


def test_select_prompt_keeps_api_response_chunk() -> None:
    q = "What are the success response fields for the QR generation API?"
    intents = detect_query_intents(q)
    hdr = (
        "API Reference ID: API-REST-QRU-01\n"
        "Service Name: QR Checkout URL Endpoint\n"
        "Section: API Metadata\n"
        "Service Method: POST\n"
    )
    meta_txt = hdr + "Service Type: REST (JSON)\n"
    resp_txt = (
        "API Reference ID: API-REST-QRU-01\n"
        "Service Name: QR Checkout URL Endpoint\n"
        "Section: API Response Parameters\n"
        "Response Parameters: status (string, Mandatory): ok; code (string, Mandatory): x"
    )
    results = [
        {"chunk_type": "api_metadata_chunk", "chunk_text": meta_txt, "api_reference_id": "API-REST-QRU-01", "document_id": 1},
        {"chunk_type": "api_overview_chunk", "chunk_text": hdr + "Overview body.\n", "api_reference_id": "API-REST-QRU-01", "document_id": 1},
        {"chunk_type": "api_response_parameters_chunk", "chunk_text": resp_txt, "api_reference_id": "API-REST-QRU-01", "document_id": 1},
    ]
    rs = RagService()
    selected, _diag = rs._select_prompt_contexts(
        results=results, top_k=5, detected_intents=intents, question=q
    )
    types = [c.get("chunk_type") for c in selected]
    assert "api_response_parameters_chunk" in types
    joined = "\n".join((c.get("chunk_text") or "") for c in selected)
    assert "status" in joined and "Response Parameters" in joined


def test_suggested_questions_use_response_chunk_scope() -> None:
    contexts = [
        {"chunk_type": "generic_section_chunk", "service_name": "deactivationService", "api_reference_id": "API-REST-OLD-01"},
        {"chunk_type": "api_response_parameters_chunk", "service_name": "QR Checkout URL Endpoint", "api_reference_id": "API-REST-QRU-01"},
    ]
    out = suggested_question_service.generate(
        user_question="What are the success response fields for the QR generation API?",
        answer="The documented response includes status, code, desc, timestamp, and related fields.",
        contexts=contexts,
        document_type="api",
        detected_intents=["response_field_intent", "parameter_intent"],
    )
    assert any("QR Checkout URL Endpoint" in s for s in out)
    assert not any("deactivationService" in s for s in out)
